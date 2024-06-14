from docx import Document
import re
from docx.shared import Inches
import os
import random
import sys
from art import text2art
from datetime import date
import json
import xml.etree.ElementTree as ET

#文件特殊符号检查
def file_check(filename):
    with open(filename, 'r+') as file:
        lines = file.readline()
        xml_content = file.read()
        if '<!-- check -->' in xml_content:
            file.close()
        else:
            print('文件:'+filename+'未处理')
            xml_content = xml_content.replace('&', '&amp;')
            xml_content = re.sub(r'(<vuln-descript>.*?<\/vuln-descript>)',lambda match: match.group(1).replace('<', '&lt;').replace('>','&gt;'), xml_content, flags=re.DOTALL)
            xml_content = xml_content.replace('&lt;vuln-descript&gt;','<vuln-descript>').replace('&lt;/vuln-descript&gt;','</vuln-descript>')
            with open(filename, 'w') as file:
                  file.write('<?xml version="1.0" encoding="UTF-8"?>\n<!-- check -->\n')
                  file.write(xml_content)
            file.close()
            print('文件:'+filename+'处理完毕')
#解析dependency-check输出的json报告
def dependency_check(filename):
    with open(filename, 'r') as file:
        data = json.load(file)
        file.close()
    vulnerabilities = []
    for dependency in data['dependencies']:
        if 'vulnerabilities' in dependency:
            for vulnerability in dependency['vulnerabilities']:
                vulnerabilities.append({
                    'jar': dependency['fileName'],
                    'cve': vulnerability['name'],
                    'path': dependency['filePath'],
                })
    return vulnerabilities
#获取cnnvd信息并生成报告
def CNNVD_db(cve):
    parts = re.split('-', cve)
    files = os.getcwd() + '/cnnvd_xml'
    for file in os.listdir(files):
         xmll = re.search("^....",file).group()
         if xmll == parts[1]:
            # print(file)
             tree = ET.parse(files + '/' + file)
             root = tree.getroot()
             for entry in root.findall('.//entry'):
                cveid = entry.find('.//cve-id').text.strip()
                if cveid == cve:
                    cnnvd = []
                    cnnvd.append(entry.find('.//name').text.strip())
                    cnnvd.append(entry.find('.//severity').text.strip())
                    cnnvd.append(entry.find('.//vuln-descript').text.strip())
                    return cnnvd
#生成报告
def word_result(jsonname):
    vulnerabilities = dependency_check(jsonname)
    doc = Document()
    doc.add_heading('dependency-check检查报告', level=1)
    #doc.add_picture('pie_chart.png', width=Inches(4))
    jar_path_dict = {}
    for vuln in vulnerabilities:
        jar = vuln['jar']
        path = vuln['path']
        cve_id = vuln['cve']
        if jar not in jar_path_dict:
            jar_path_dict[jar] = path
            doc.add_heading(jar, level=2)
            doc.add_paragraph('jar包路径: ' + path)
        doc.add_heading(cve_id, level=3)
        cnnvd_result = CNNVD_db(cve_id)
        doc.add_paragraph('漏洞名称: ' + cnnvd_result[0])
        doc.add_paragraph('漏洞等级: ' + cnnvd_result[1])
        doc.add_paragraph('漏洞描述: ' + cnnvd_result[2])
    doc.save(os.getcwd() + '/' + date.today().strftime('%Y-%m-%d') + '-' + generate_random_characters(5) + '_vuln_report.docx')

def generate_random_characters(length):
    chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789'
    return ''.join(random.choices(chars, k=length))

if __name__ == '__main__':
    banner_text = text2art("0X00001")
    print(banner_text)
    print("将dependency-check报告解析为cnnvd版本版本，只能输出word报告")
    print("Parse the dependency check report to the CNNVD version, only able to output Word reports")
    if len(sys.argv) != 2:
        print("Usage: python cnnvd.py <dependency-check_json_file_path>")
    else:
        jsonname = sys.argv[1]
        try:
            word_result(jsonname)
            print("大侠，报告转换成功！")
        except Exception as e:
            print("报错：请输入正确的json路径")